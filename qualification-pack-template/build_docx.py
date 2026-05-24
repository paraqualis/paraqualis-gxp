#!/usr/bin/env python3
"""build_docx.py — render every docs/*.md into a branded ParaQualis .docx.

Markdown is the source of truth; the .docx is a generated rendering. Re-run anytime:
    python3 build_docx.py
Handles: # headings, | tables |, ``` code fences, > blockquotes, - bullets, **bold**.
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DEEP_BLUE = RGBColor(0x27, 0x6B, 0xB2)   # ParaQualis primary
SLATE = RGBColor(0x81, 0xA5, 0xC1)
HDR_FILL = "276BB2"
DOCS = Path(__file__).parent / "docs"


def _shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hex_fill)
    tcPr.append(sh)


def _bold_runs(paragraph, text):
    # split on **bold** and add runs
    for i, seg in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if not seg:
            continue
        run = paragraph.add_run(seg)
        if i % 2 == 1:
            run.bold = True


def _add_table(doc, rows):
    cols = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=cols)
    t.style = "Light Grid Accent 1"
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci in range(cols):
            txt = row[ci] if ci < len(row) else ""
            cell = cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            if ri == 0:
                _shade(cell, HDR_FILL)
                run = p.add_run(re.sub(r"\*\*", "", txt))
                run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(8.5)
            else:
                _bold_runs(p, txt)
                for r in p.runs:
                    r.font.size = Pt(8.5)


def render(md_path: Path, out_path: Path):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i, n = 0, len(lines)
    table_buf, code_buf, in_code = [], [], False

    def flush_table():
        if not table_buf:
            return
        rows = []
        for ln in table_buf:
            if re.match(r"^\s*\|[\s:|-]+\|\s*$", ln):   # separator row
                continue
            cells = [c.strip() for c in ln.strip().strip("|").split("|")]
            rows.append(cells)
        if rows:
            _add_table(doc, rows)
        table_buf.clear()

    while i < n:
        ln = lines[i]
        if ln.strip().startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Consolas"; run.font.size = Pt(8)
                _shade_paragraph(p)
                code_buf.clear(); in_code = False
            else:
                flush_table(); in_code = True
            i += 1; continue
        if in_code:
            code_buf.append(ln); i += 1; continue
        if ln.strip().startswith("|"):
            table_buf.append(ln); i += 1; continue
        else:
            flush_table()
        if ln.startswith("# "):
            h = doc.add_heading(level=0); r = h.add_run(ln[2:]); r.font.color.rgb = DEEP_BLUE
        elif ln.startswith("## "):
            h = doc.add_heading(level=1); r = h.add_run(ln[3:]); r.font.color.rgb = DEEP_BLUE
        elif ln.startswith("### "):
            h = doc.add_heading(level=2); r = h.add_run(ln[4:]); r.font.color.rgb = DEEP_BLUE
        elif ln.strip().startswith("> "):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
            r = p.add_run(ln.strip()[2:]); r.italic = True; r.font.color.rgb = SLATE
        elif re.match(r"^\s*[-*] ", ln):
            p = doc.add_paragraph(style="List Bullet"); _bold_runs(p, re.sub(r"^\s*[-*] ", "", ln))
        elif ln.strip() == "":
            pass
        else:
            p = doc.add_paragraph(); _bold_runs(p, ln)
        i += 1
    flush_table()
    doc.save(str(out_path))
    print(f"  wrote {out_path.name}")


def _shade_paragraph(p):
    pPr = p._p.get_or_add_pPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), "EFEAE3")  # warm cream
    pPr.append(sh)


if __name__ == "__main__":
    mds = sorted(DOCS.glob("*.md"))
    if not mds:
        sys.exit("no docs/*.md found")
    print("Building branded .docx:")
    for md in mds:
        render(md, md.with_suffix(".docx"))
    print("Done.")
