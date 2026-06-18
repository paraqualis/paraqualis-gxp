#!/usr/bin/env python3
"""build_docx.py — render every docs/*.md into a branded ParaQualis .docx.

Markdown is the source of truth; the .docx is a generated rendering. Re-run anytime:
    python3 build_docx.py
Landscape, with an explicit consistent type scale (no Word built-in heading styles,
which vary by template and cause sizes to jump). Handles: # headings, | tables |,
``` code fences, > blockquotes, - bullets, **bold**, --- rules.
"""
import re
import sys
from pathlib import Path

# Pre-flight: fail loud with guidance if the rendering dependency is absent,
# rather than dying on an opaque ImportError mid-run.
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError as e:
    sys.exit(
        f"build_docx.py: missing required dependency ({e.name}). "
        "This script needs 'python-docx'.\n"
        "Install the pinned dependencies first:  pip install -r requirements.txt\n"
        "(or: pip install python-docx)"
    )

DEEP_BLUE = RGBColor(0x27, 0x6B, 0xB2)   # ParaQualis primary
SLATE = RGBColor(0x81, 0xA5, 0xC1)
HDR_FILL = "276BB2"
CODE_FILL = "EFEAE3"

# One explicit type scale — applied to every run, so nothing inherits a stray size.
BODY_PT = 10
TABLE_PT = 9
CODE_PT = 8.5
H_PT = {0: 18, 1: 14, 2: 12, 3: 11}

DOCS = Path(__file__).parent / "docs"


def _shade_cell(cell, hex_fill):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(sh)


def _shade_para(p, hex_fill):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hex_fill)
    p._p.get_or_add_pPr().append(sh)


def _runs(paragraph, text, size_pt, color=None):
    """Add text to a paragraph, honouring **bold**, every run at size_pt."""
    for i, seg in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if seg == "":
            continue
        r = paragraph.add_run(seg)
        r.font.size = Pt(size_pt)
        if color is not None:
            r.font.color.rgb = color
        if i % 2 == 1:
            r.bold = True


def _heading(doc, text, level):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10 if level <= 1 else 8)
    pf.space_after = Pt(4)
    pf.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(H_PT.get(level, 12))
    r.font.color.rgb = DEEP_BLUE


def _add_table(doc, rows):
    cols = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=cols)
    t.style = "Light Grid Accent 1"
    t.autofit = True
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci in range(cols):
            txt = row[ci] if ci < len(row) else ""
            cell = cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            if ri == 0:
                _shade_cell(cell, HDR_FILL)
                r = p.add_run(re.sub(r"\*\*", "", txt))
                r.bold = True
                r.font.size = Pt(TABLE_PT)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            else:
                _runs(p, txt, TABLE_PT)


def render(md_path: Path, out_path: Path):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(BODY_PT)

    cp = doc.core_properties
    cp.author = "ParaQualis LLC"
    cp.last_modified_by = "ParaQualis LLC"
    cp.title = md_path.stem.replace("-", " ")

    # Landscape + comfortable margins (wide test-case tables need the room).
    sec = doc.sections[0]
    w, h = sec.page_width, sec.page_height
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = max(w, h), min(w, h)
    sec.left_margin = sec.right_margin = Inches(0.6)
    sec.top_margin = sec.bottom_margin = Inches(0.6)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i, n = 0, len(lines)
    table_buf, code_buf, in_code = [], [], False

    def flush_table():
        if not table_buf:
            return
        rows = []
        for ln in table_buf:
            if re.match(r"^\s*\|[\s:|-]+\|\s*$", ln):     # separator row
                continue
            rows.append([c.strip() for c in ln.strip().strip("|").split("|")])
        if rows:
            _add_table(doc, rows)
        table_buf.clear()

    while i < n:
        ln = lines[i]
        if ln.strip().startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                r = p.add_run("\n".join(code_buf))
                r.font.name = "Consolas"; r.font.size = Pt(CODE_PT)
                _shade_para(p, CODE_FILL)
                code_buf.clear(); in_code = False
            else:
                flush_table(); in_code = True
            i += 1; continue
        if in_code:
            code_buf.append(ln); i += 1; continue
        if ln.strip().startswith("|"):
            table_buf.append(ln); i += 1; continue
        flush_table()

        stripped = ln.strip()
        if stripped.startswith("<!--") and stripped.endswith("-->"):
            i += 1; continue                              # HTML comment / marker → skip
        if stripped in ("---", "***", "___"):
            i += 1; continue                              # horizontal rule → skip
        if ln.startswith("#### "):
            _heading(doc, ln[5:], 3)
        elif ln.startswith("### "):
            _heading(doc, ln[4:], 2)
        elif ln.startswith("## "):
            _heading(doc, ln[3:], 1)
        elif ln.startswith("# "):
            _heading(doc, ln[2:], 0)
        elif stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            _runs(p, stripped[2:], BODY_PT, color=SLATE)
        elif re.match(r"^\s*[-*] ", ln):
            p = doc.add_paragraph(style="List Bullet")
            _runs(p, re.sub(r"^\s*[-*] ", "", ln), BODY_PT)
        elif stripped == "":
            pass
        else:
            _runs(doc.add_paragraph(), ln, BODY_PT)
        i += 1
    flush_table()
    doc.save(str(out_path))
    print(f"  wrote {out_path.name}")


if __name__ == "__main__":
    mds = sorted(DOCS.glob("*.md"))
    if not mds:
        sys.exit("no docs/*.md found")
    print("Building branded .docx (landscape):")
    for md in mds:
        render(md, md.with_suffix(".docx"))
    print("Done.")
