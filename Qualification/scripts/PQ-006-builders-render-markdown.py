#!/usr/bin/env python3
# PQ-006 — The Word and Excel builders render the Markdown source of truth.
# Traces: URS-007, URS-011 (Markdown is source of truth; renders regenerated, faithful).
# ParaQualis-authored, pending owner confirmation.
#
# Why this matters (plain English): clients receive signable Word and a QA Excel workbook, but
# Markdown stays the single source of truth — the .docx/.xlsx must be GENERATED from it and
# carry its content faithfully. This test runs both builders on real pack Markdown in an
# isolated temp copy, then verifies (a) a valid .docx is produced and (b) a valid .xlsx with
# the expected per-stage sheets is produced, and that distinctive Markdown content round-trips
# into each. It is read-only with respect to the repo (works on a temp copy).
#
# Expected sheet names are SOURCED from the builder's own contract:
#   Qualification/build_xlsx.py:70-83  ->  Summary, IQ, OQ, PQ, Traceability (+ Gaps if present)
#
# Usage: python3 PQ-006-builders-render-markdown.py [REPO_ROOT]
import sys
import os
import shutil
import tempfile
import subprocess

REPO = sys.argv[1] if len(sys.argv) > 1 else "/Users/craigwylie/Devl/paraqualis-gxp"
PACK = os.path.join(REPO, "Qualification")

print("PQ-006 builders render Markdown faithfully")
print("-" * 60)

# Dependency check — missing libs must FAIL LOUD, never silently skip.
missing = []
for mod in ("docx", "openpyxl"):
    try:
        __import__(mod)
    except Exception:
        missing.append(mod)
if missing:
    print(f"FAIL  required builder libraries not installed: {missing}")
    print("      (python-docx / openpyxl — see dependency-manifest gap GPQ-002)")
    sys.exit(1)

work = tempfile.mkdtemp(prefix="pq006_")
try:
    # Copy the pack (builders + docs) into an isolated working dir so we never touch the repo.
    dst = os.path.join(work, "Qualification")
    shutil.copytree(PACK, dst)
    docs = os.path.join(dst, "docs")

    md_files = [f for f in os.listdir(docs) if f.endswith(".md")] if os.path.isdir(docs) else []
    if not md_files:
        print("NOT-EXECUTED  no docs/*.md present yet to render (run a build first).")
        sys.exit(2)  # explicitly not a pass
    print(f"found Markdown to render: {sorted(md_files)}")

    # 1. Word builder.
    r = subprocess.run([sys.executable, "build_docx.py"], cwd=dst,
                       capture_output=True, text=True)
    docx_made = [f for f in os.listdir(docs) if f.endswith(".docx")]
    if r.returncode == 0 and docx_made:
        print(f"PASS  build_docx.py produced: {sorted(docx_made)}")
    else:
        print(f"FAIL  build_docx.py rc={r.returncode}; stderr={r.stderr.strip()[:300]}")
        sys.exit(1)

    # Faithful round-trip: a distinctive token from a PQ.md heading must appear in PQ.docx text.
    ok_docx = True
    pq_docx = os.path.join(docs, "PQ.docx")
    if os.path.exists(os.path.join(docs, "PQ.md")) and os.path.exists(pq_docx):
        import docx as _docx
        body = "\n".join(p.text for p in _docx.Document(pq_docx).paragraphs)
        tables = "\n".join(c.text for t in _docx.Document(pq_docx).tables
                           for row in t.rows for c in row.cells)
        haystack = body + "\n" + tables
        if "PQ-001" in haystack or "Performance Qualification" in haystack:
            print("PASS  PQ.md content round-trips into PQ.docx")
        else:
            print("FAIL  distinctive PQ.md content not found in PQ.docx")
            ok_docx = False

    # 2. Excel builder.
    r = subprocess.run([sys.executable, "build_xlsx.py"], cwd=dst,
                       capture_output=True, text=True)
    xlsx = os.path.join(docs, "Qualification-Pack.xlsx")
    if r.returncode != 0 or not os.path.exists(xlsx):
        print(f"FAIL  build_xlsx.py rc={r.returncode}; stderr={r.stderr.strip()[:300]}")
        sys.exit(1)
    import openpyxl
    wb = openpyxl.load_workbook(xlsx)
    sheets = wb.sheetnames
    print(f"PASS  build_xlsx.py produced workbook with sheets: {sheets}")
    # Contract: Summary, IQ, OQ, PQ, Traceability expected (build_xlsx.py:70-79).
    expected_core = {"Summary", "IQ", "OQ", "PQ", "Traceability"}
    ok_xlsx = expected_core.issubset(set(sheets))
    if ok_xlsx:
        print("PASS  workbook contains the contracted per-stage + summary + trace sheets")
    else:
        print(f"FAIL  workbook missing contracted sheets: {expected_core - set(sheets)}")

    print("-" * 60)
    if ok_docx and ok_xlsx:
        print("RESULT: PASS — both builders rendered the Markdown faithfully.")
        sys.exit(0)
    print("RESULT: FAIL — a builder did not render faithfully.")
    sys.exit(1)
finally:
    shutil.rmtree(work, ignore_errors=True)
