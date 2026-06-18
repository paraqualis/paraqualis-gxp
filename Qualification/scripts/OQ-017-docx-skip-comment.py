#!/usr/bin/env python3
#
# Copyright (c) 2026 ParaQualis LLC
# Licensed under the MIT License — see LICENSE in the repository root.
#
"""OQ-017 — build_docx.render() skips HTML comment / PARAQUALIS-LOCK marker lines.

Requirement (OQ.md OQ-017): lines matching <!-- ... --> (HTML comments, including the
PARAQUALIS-LOCK marker) are silently skipped during rendering and must not appear in the
document body.

Behaviour SOURCED from qualification-pack-template/build_docx.py:146-147:
    if stripped.startswith("<!--") and stripped.endswith("-->"):
        i += 1; continue                              # HTML comment / marker -> skip

Read-only / safe: renders into a private temp dir, then RE-OPENS the .docx and asserts the
marker text is absent from every paragraph and table cell. Requires python-docx; FAILS LOUD
if absent. PASS/FAIL emitted; exits non-zero on FAIL.
"""
import importlib.util
import os
import sys
import tempfile
from pathlib import Path

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
BUILD_DOCX = os.path.join(ROOT, "qualification-pack-template", "build_docx.py")
MARKER = "PARAQUALIS-LOCK"

try:
    spec = importlib.util.spec_from_file_location("build_docx", BUILD_DOCX)
    build_docx = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(build_docx)
    from docx import Document
except ModuleNotFoundError as e:
    print(f"FAIL: required dependency missing for build_docx ({e}). Run: pip install python-docx")
    sys.exit(1)

MD = """# Approved Document

<!-- PARAQUALIS-LOCK: approved -->

Visible body paragraph.

<!-- a plain html comment -->

More visible text.
"""


def all_text(doc):
    chunks = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    return "\n".join(chunks)


with tempfile.TemporaryDirectory() as d:
    md_path = Path(d) / "approved.md"
    out_path = Path(d) / "approved.docx"
    md_path.write_text(MD, encoding="utf-8")
    try:
        build_docx.render(md_path, out_path)
    except Exception as e:
        print(f"FAIL: render() raised an exception: {e!r}")
        sys.exit(1)

    text = all_text(Document(str(out_path)))
    if MARKER in text or "<!--" in text:
        print(f"FAIL: comment/marker leaked into document body (found {MARKER!r} or '<!--')")
        sys.exit(1)
    if "Visible body paragraph." in text:
        print("PASS: comment/marker lines skipped; visible content retained")
        sys.exit(0)
    print("FAIL: expected visible content missing — render may have dropped real body text")
    sys.exit(1)
